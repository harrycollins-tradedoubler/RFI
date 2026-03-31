#!/usr/bin/env python3
"""
Build a single Ready-only RAG dataset CSV from mixed document inputs.

Supported inputs:
- PDF (.pdf)
- Word (.docx)
- Excel (.xlsx)
- Plain text-like files (.txt, .md, .csv, .json, .log, .rst)
"""

from __future__ import annotations

import argparse
import csv
import json
import re
from pathlib import Path
from typing import Iterable

from openpyxl import load_workbook
from pypdf import PdfReader

try:
    from docx import Document
except Exception:  # pragma: no cover - dependency guard
    Document = None


TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".log", ".rst"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | {".pdf", ".docx", ".xlsx"}


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.strip() for line in text.split("\n")]
    lines = [line for line in lines if line]
    merged = " ".join(lines)
    return re.sub(r"\s+", " ", merged).strip()


def chunk_text(text: str, chunk_size: int, overlap: int) -> list[str]:
    if not text:
        return []

    chunks: list[str] = []
    start = 0
    n = len(text)
    while start < n:
        end = min(n, start + chunk_size)
        if end < n:
            floor = start + int(chunk_size * 0.65)
            para_break = text.rfind(". ", floor, end)
            if para_break > start:
                end = para_break + 1

        piece = text[start:end].strip()
        if piece:
            chunks.append(piece)

        if end >= n:
            break
        start = max(end - overlap, start + 1)

    return chunks


def discover_files(inputs: list[str]) -> list[Path]:
    paths: list[Path] = []
    for raw in inputs:
        p = Path(raw).resolve()
        if not p.exists():
            continue
        if p.is_file():
            if p.suffix.lower() in SUPPORTED_EXTENSIONS:
                paths.append(p)
            continue
        for f in p.rglob("*"):
            if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS:
                paths.append(f.resolve())

    # preserve deterministic order + unique paths
    unique = []
    seen = set()
    for p in sorted(paths, key=lambda x: str(x).lower()):
        key = str(p).lower()
        if key not in seen:
            seen.add(key)
            unique.append(p)
    return unique


def extract_pdf_units(path: Path) -> list[dict]:
    reader = PdfReader(str(path))
    units = []
    for page_num, page in enumerate(reader.pages, start=1):
        txt = normalize_text(page.extract_text() or "")
        if txt:
            units.append({"unit_label": f"Page {page_num}", "text": txt, "page": page_num})
    return units


def extract_docx_units(path: Path) -> list[dict]:
    if Document is None:
        raise RuntimeError("python-docx not available.")

    doc = Document(str(path))
    lines = []

    for p in doc.paragraphs:
        t = normalize_text(p.text or "")
        if t:
            lines.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [normalize_text(c.text or "") for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                lines.append(" | ".join(cells))

    merged = normalize_text("\n".join(lines))
    if not merged:
        return []
    return [{"unit_label": "Document", "text": merged}]


def extract_xlsx_units(path: Path) -> list[dict]:
    wb = load_workbook(str(path), data_only=True, read_only=True)
    units = []
    for ws in wb.worksheets:
        lines = []
        for row in ws.iter_rows(values_only=True):
            vals = ["" if v is None else normalize_text(str(v)) for v in row]
            vals = [v for v in vals if v]
            if vals:
                lines.append(" | ".join(vals))
        merged = normalize_text("\n".join(lines))
        if merged:
            units.append({"unit_label": f"Sheet {ws.title}", "text": merged, "sheet": ws.title})
    return units


def extract_text_units(path: Path) -> list[dict]:
    txt = path.read_text(encoding="utf-8", errors="ignore")
    txt = normalize_text(txt)
    if not txt:
        return []
    return [{"unit_label": "Text", "text": txt}]


def extract_units(path: Path) -> list[dict]:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf_units(path)
    if ext == ".docx":
        return extract_docx_units(path)
    if ext == ".xlsx":
        return extract_xlsx_units(path)
    if ext in TEXT_EXTENSIONS:
        return extract_text_units(path)
    return []


def iter_rows(
    files: list[Path],
    chunk_size: int,
    overlap: int,
    id_prefix: str,
) -> Iterable[dict]:
    rec_id = 1
    for path in files:
        ext = path.suffix.lower().lstrip(".")
        units = extract_units(path)
        if not units:
            continue

        for u_idx, unit in enumerate(units, start=1):
            chunks = chunk_text(unit["text"], chunk_size=chunk_size, overlap=overlap)
            for c_idx, chunk in enumerate(chunks, start=1):
                qid = f"{id_prefix}-{rec_id:05d}"
                metadata = {
                    "source_file": path.name,
                    "source_path": str(path),
                    "source_type": ext,
                    "unit_label": unit.get("unit_label", ""),
                    "unit_index": u_idx,
                    "chunk_index": c_idx,
                    "chunks_in_unit": len(chunks),
                }
                if "page" in unit:
                    metadata["page"] = unit["page"]
                if "sheet" in unit:
                    metadata["sheet"] = unit["sheet"]

                rel_parent = path.parent.name
                yield {
                    "question_id": qid,
                    "content": chunk,
                    "canonical_question": "",
                    "starter_answer": "",
                    "category": "Attached Source Corpus",
                    "subcategory": f"{ext.upper()} | {rel_parent}",
                    "response_type": "document_chunk",
                    "item_type": "chunk",
                    "language": "en",
                    "evidence_reference": f"{path.name} - {unit.get('unit_label', 'Document')}",
                    "source_files": path.name,
                    "source_sections": unit.get("unit_label", "Document"),
                    "readiness_status": "Ready",
                    "metadata_json": json.dumps(metadata, ensure_ascii=False),
                }
                rec_id += 1


def write_dataset(rows: list[dict], output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    columns = [
        "question_id",
        "content",
        "canonical_question",
        "starter_answer",
        "category",
        "subcategory",
        "response_type",
        "item_type",
        "language",
        "evidence_reference",
        "source_files",
        "source_sections",
        "readiness_status",
        "metadata_json",
    ]
    with output.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=columns)
        w.writeheader()
        w.writerows(rows)


def main() -> int:
    parser = argparse.ArgumentParser(description="Build combined Ready dataset from attachments.")
    parser.add_argument(
        "--input",
        action="append",
        required=True,
        help="Input file or directory. Repeat this argument for multiple roots.",
    )
    parser.add_argument(
        "--output",
        default="output/dataset/attached_sources_ready_dataset.csv",
        help="Output CSV path.",
    )
    parser.add_argument("--chunk-size", type=int, default=1100)
    parser.add_argument("--overlap", type=int, default=180)
    parser.add_argument("--id-prefix", default="ATTACH")
    args = parser.parse_args()

    files = discover_files(args.input)
    if not files:
        raise SystemExit("No supported files found in inputs.")

    rows = list(
        iter_rows(
            files=files,
            chunk_size=args.chunk_size,
            overlap=args.overlap,
            id_prefix=args.id_prefix,
        )
    )
    if not rows:
        raise SystemExit("No extractable text found from selected files.")

    output = Path(args.output).resolve()
    write_dataset(rows, output)

    print(f"Files processed: {len(files)}")
    print(f"Rows generated: {len(rows)}")
    print(f"Output: {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
